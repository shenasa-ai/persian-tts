# from docx import Document
# import re
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
path = os.getcwd()
direct = os.listdir()

for i in direct:
#     document = Document()
#     document.add_heading(i, 0)
#     myfile = open(i,'r',encoding='utf-8')
#     myfile = myfile.read()
#     myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
#     p = document.add_paragraph(myfile)
#     document.save(i+'.docx')

 
    doc = Document()
    with open(i, 'r', encoding='utf-8') as file:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = par.add_run(file.read())
        font = run.font
        font.rtl = True
    doc.save(i[0:-4]+".docx")
